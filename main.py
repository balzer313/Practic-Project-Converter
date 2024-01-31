import tkinter as tk
from tkinter import ttk, filedialog
import json
import openpyxl
import openpyxl.styles
from openpyxl.utils import get_column_letter
import pandas as pd
from PIL import Image, ImageTk
from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import filedialog
import pdfplumber


def pdf_convert(invoice_file,output_excel):
    global money_type
    if money_type == None: money_type = 'DOLLAR'
    def add_discounts_to_line(dict_line):
        for discount in dict_line['discounts']:
            reason = discount['reason']
            amount = discount['amount']
            if amount != 0 and amount:
                percentage = discount['percentage']
                if percentage and percentage != 0:
                    dict_line["{}_percentage".format(reason)] = percentage
        return dict_line

    def order_keys(dict_line, proper_order):
        new_dict_line = dict()
        for element in proper_order:
            new_dict_line[element] = dict_line[element]
        return new_dict_line

    def do_calculations(line):
        if 'loyaltyDiscountTotal' not in line.keys():
            line['loyaltyDiscountTotal'] = 0
            line['loyaltyDiscountTotal_percentage'] = 0
        else:
            if line['loyaltyDiscountTotal'] != 0:
                line['loyaltyDiscountTotal_percentage'] = round(line['loyaltyDiscountTotal'] / line['unitPrice'] * 100)
            else:
                line['loyaltyDiscountTotal_percentage'] = 0
        if 'upgradeCredit' not in line.keys():
            line['upgradeCredit'] = 0
        if 'priceAdjustment' not in line.keys():
            line['priceAdjustment'] = 0
        line['priceAdjustment'] = line['priceAdjustment'] - line['partnerDiscountTotal'] - line['loyaltyDiscountTotal']
        discounted = line['unitPrice'] - line['loyaltyDiscountTotal'] - line['upgradeCredit'] - line['priceAdjustment']
        if 'partnerDiscountTotal' not in line.keys():
            line['partnerDiscountTotal'] = 0
        else:
            if line['partnerDiscountTotal'] != 0:
                line['partnerDiscountTotal_percentage'] = round(line['partnerDiscountTotal'] / discounted * 100)
            else:
                line['partnerDiscountTotal_percentage'] = 0
        return line

    def rename_dict_keys(full_data):
        rename_matrix = {
            'productName': 'Product name',
            'unitPrice': 'Unit price',
            'loyaltyDiscountTotal': 'Loyalty discount',
            'upgradeCredit': 'Upgrade credit',
            'priceAdjustment': 'Price adjustment',
            'partnerDiscountTotal': 'Partner discount',
            'total': 'Total price for partner',
            'loyaltyDiscountTotal_percentage': 'Loyalty discount %',
            'partnerDiscountTotal_percentage': 'Partner discount %'
        }
        renamed_dict = full_data
        for key in rename_matrix.keys():
            value = rename_matrix[key]
            renamed_dict[value] = full_data[key]
            del renamed_dict[key]
        return renamed_dict

    def format_products(invoice_data):
        formatted = []
        ordered_columns_to_keep = ['productName', 'unitPrice', 'loyaltyDiscountTotal', 'upgradeCredit',
                                   'priceAdjustment',
                                   'partnerDiscountTotal', 'total', 'loyaltyDiscountTotal_percentage',
                                   'partnerDiscountTotal_percentage']
        for line in invoice_data:
            new_line = do_calculations(line)
            line['productName'] = "{}, {} users".format(line['productName'], line['unitCount'])
            line_keys = list(new_line.keys())
            for element in line_keys:
                if element not in ordered_columns_to_keep:
                    del new_line[element]
            new_line = order_keys(new_line, ordered_columns_to_keep)
            if new_line['unitPrice'] != 0:
                new_line = rename_dict_keys(new_line)
                formatted.append(new_line)
        return formatted

    def create_excel_header(company_name, invoice_number):
        header = (
            ('Company name:', company_name),
            ('Quote/Invoice number:', invoice_number),
            ('Discount for customer:', '5'),
            ()
        )
        return header

    def add_header_lines_and_save(header, excel_file_name, new_sheet_name="Main data"):
        wb = openpyxl.load_workbook(excel_file_name)
        ws = wb.active
        sheet_name = wb.sheetnames[0]
        new_ws = wb.create_sheet(new_sheet_name)
        header_lines = len(header)
        data_lines = 0

        for row in header:
            new_ws.append(row)
        for row in ws.iter_rows(values_only=True):
            new_ws.append(row)
            data_lines = data_lines + 1
        start_sum = header_lines + 1
        end_sum = header_lines + data_lines
        row_to_add = (
            "Total",
            "=SUM(B{}:B{})".format(start_sum, end_sum),
            "=SUM(C{}:C{})".format(start_sum, end_sum),
            "=SUM(D{}:D{})".format(start_sum, end_sum),
            "=SUM(E{}:E{})".format(start_sum, end_sum),
            "=SUM(F{}:F{})".format(start_sum, end_sum),
            "=SUM(G{}:G{})".format(start_sum, end_sum),
            "",
            "",
            "=SUM(J{}:J{})".format(start_sum, end_sum)
        )
        new_ws.append(row_to_add)
        del wb[sheet_name]
        wb.save(excel_file_name)

    def add_column_for_customer_price(table, header_size, percentage_cell):
        table_with_new_column = table
        for row in table:
            line_number = table.index(row) + header_size + 2
            row['Customer price'] = "=B{}*(100-{})/100".format(line_number, percentage_cell)
        return table_with_new_column

    def add_totals_format(file_name, number_line):
        wb = openpyxl.load_workbook(file_name)
        ws = wb.active
        cells = ["A{}".format(number_line),
                 "B{}".format(number_line),
                 'C{}'.format(number_line),
                 'D{}'.format(number_line),
                 'E{}'.format(number_line),
                 'F{}'.format(number_line),
                 'G{}'.format(number_line),
                 'H{}'.format(number_line),
                 'I{}'.format(number_line),
                 'J{}'.format(number_line)]
        for cell in cells:
            cell_to_operate = ws[cell]
            cell_to_operate.font = openpyxl.styles.Font(bold=True)
            if cell != "A{}".format(number_line):
                cell_to_operate.fill = openpyxl.styles.PatternFill(start_color='FFF2CC', end_color='FFF2CC',
                                                                   fill_type='solid')
        wb.save(file_name)

    def add_global_sheet_formatting(file_name, header_size, data_size):
        wb = openpyxl.load_workbook(file_name)
        ws = wb.active
        for column_cells in ws.columns:
            new_column_length = max(len(str(cell.value)) for cell in column_cells)
            new_column_letter = (get_column_letter(column_cells[0].column))
            if new_column_length > 0:
                ws.column_dimensions[new_column_letter].width = new_column_length * 1.23
        columns = ["B", "C", "D", "E", "F", "G", "H", "I", "J"]
        line_to_work = header_size + 2
        while line_to_work <= header_size + data_size + 2:
            for column in columns:
                cell_key_to_work = "{}{}".format(column, line_to_work)
                cell = ws[cell_key_to_work]
                if column != "H" and column != "I":
                    cell.number_format = '# ##0.00'
                else:
                    if cell.value:
                        cell.value = cell.value / 100
                        cell.number_format = '0%'
            line_to_work = line_to_work + 1
        wb.save(file_name)

    new_sheet = "Main data"
    with open(invoice_file) as f:
        data = json.load(f)
    products = data['orderItems']
    invoice_number = data['orderNumber']
    company_name = data['orderItems'][0]['licensedTo']
    header = create_excel_header(company_name=company_name,
                                 invoice_number=invoice_number)
    formatted_products = format_products(products)
    header_lines = len(header)
    products_lines = len(formatted_products)
    formatted_products = add_column_for_customer_price(table=formatted_products,
                                                       header_size=header_lines,
                                                       percentage_cell="B3")
    df = pd.DataFrame(formatted_products)
    def multiply_unit_price_dirham(row):
        return row['Unit price'] * 3.67
    if money_type == 'DIRHAM': df['Unit price'] = df.apply(multiply_unit_price_dirham, axis=1)

    df.to_excel(output_excel, index=False)
    add_header_lines_and_save(header=header,
                              excel_file_name=output_excel,
                              new_sheet_name=new_sheet)
    add_totals_format(file_name=output_excel,
                      number_line=header_lines + 1)
    add_totals_format(file_name=output_excel,
                      number_line=header_lines + products_lines + 2)
    add_global_sheet_formatting(file_name=output_excel,
                                header_size=header_lines,
                                data_size=products_lines)


def convert():
    global selected_json_file
    global money_type
    if selected_json_file:
        output_folder = filedialog.askdirectory(title="Select Folder to Save File")
        if output_folder:
            print('converting...')
            pdf_convert(selected_json_file, f'{output_folder}/excel_output.xlsx')
            excel_file = f'{output_folder}/excel_output.xlsx'
            df = pd.read_excel(excel_file, header=None)  # read the excel

            df = df.fillna('')  # fill the blanks with ''
            data_dict = df.T.to_dict(orient='list')  # make dictionary of data

            company_name = ''
            table_dict = {}

            for index, row in enumerate(data_dict.values()):
                for list_index, item in enumerate(row):
                    if item == 'Company name:' and index == 0:
                        try:
                            company_name = row[list_index + 1]
                            # print(f'the company name is: {company_name}')
                        except IndexError:
                            pass
                    else:
                        if index >= 5 and item != '':
                            if list_index == 0: table_dict[item] = row[list_index + 1]

            def find_total(table_dict):  # find total
                total = 0.0
                for num in table_dict.values():
                    try:
                        total += num
                    except Exception: pass
                return total

            # print(company_name)
            # print(table_dict)
            total_of_dict = find_total(table_dict)
            # print(total_of_dict)

            # add total to table_dict
            table_dict['Total'] = total_of_dict
            table_dict['Total Due After Discount'] = 0.0

            def change_word(docs, old_word, new_word, bold, size):
                for paragraph in docs.paragraphs:
                    if f'%%{old_word}%%' in paragraph.text:
                        # print('ok')
                        paragraph.text = paragraph.text.replace(f'%%{old_word}%%',
                                                                new_word)  # Replace old_word with the new_word
                        for run in paragraph.runs:
                            if new_word in run.text:
                                run.font.size = Pt(size)  # Set font size to 18pt
                                run.font.bold = bold  # Set text to bold
                                run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the text

            if money_type == 'DIRHAM': doc = Document('documents/dubai_word.docx')
            else: doc = Document('documents/word.docx')

            change_word(doc, 'name', company_name, True, 18)

            current_datetime = datetime.now()
            formatted_date = current_datetime.strftime('%d %B %Y')
            change_word(doc, 'date', formatted_date, False, 18)

            table = doc.tables[0]
            table.style = 'Table Grid'

            table.cell(0, 0).text = 'Product'
            for cell in table.cell(0, 0).paragraphs[0].runs:
                cell.font.bold = True
            if money_type == 'DIRHAM': table.cell(0, 1).text = 'Total (AED)'
            else: table.cell(0, 1).text = 'Total (USD)'
            for cell in table.cell(0, 1).paragraphs[0].runs:
                cell.font.bold = True

            # Add data to the table
            for product, price in table_dict.items():
                row = table.add_row().cells
                row[0].text = product
                row[1].text = str(price)

            last_two_rows_indices = range(len(table_dict) - 1, len(table_dict)+1)
            # Apply bold formatting to the last two rows' product names
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if j == 0:  # Check if it's the product name column
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if i in last_two_rows_indices:
                                    run.font.bold = True

            # Save the modified document
            doc.save(f'{output_folder}/modified_existing_document.docx')
            print(f'Converting Excel to Docx and saving it to: {output_folder}')
        root.destroy()  # Close the window



def convert_json_to_docx():#select the json and move to screen2
    global selected_json_file
    selected_json_file = filedialog.askopenfilename(title="Select Json File", filetypes=(("Json Files", "*.json"), ("All Files", "*.*")))
    if not selected_json_file:
        return  # User canceled file selection
    # ######convert pdf to json
    # with pdfplumber.open(selected_pdf_file) as pdf:
    #     text_content = ''
    #     for page in pdf.pages:
    #         text_content += page.extract_text()
    # # Structure the extracted text (you may need to customize this part)
    # structured_data = text_content
    # json_data = json.dumps(structured_data, indent=2)
    # with open(selected_json_file, 'w') as json_file:
    #     json_file.write(json_data)
    # ######
    display_selected_file()

def display_selected_file():#screen 2
    global frame2  # Declare new_frame as a global variable
    frame.grid_forget()  # Remove the existing frame
    frame2 = ttk.Frame(root, padding=20)
    frame2.grid(row=0, column=0)

    new_title_label = ttk.Label(frame2, text="Selected Excel File", font=("Arial", 16, "bold"))
    new_title_label.grid(row=0, column=0, pady=(0, 20), columnspan=2)

    excel_file_label = ttk.Label(frame2, text=selected_json_file, font=("Arial", 10))
    excel_file_label.grid(row=1, column=0, pady=(0, 10), padx=10, sticky='w')

    change_excel_button = ttk.Button(frame2, text="Change Json File", command=change_selected_file)
    change_excel_button.grid(row=2, column=0, pady=(10, 10), padx=10, sticky='w')

    change_excel_button = ttk.Button(frame2, text="Convert", command=convert)
    change_excel_button.grid(row=3, column=0, pady=(10, 10), padx=10, sticky='w')

    # Center all widgets both vertically and horizontally
    frame2.grid_rowconfigure(1, weight=1)
    frame2.grid_columnconfigure(0, weight=1)

def change_selected_file():
    global selected_json_file
    global frame2  # Make new_frame accessible
    selected_json_file = None
    frame2.grid_forget()  # Remove the existing frame
    frame.grid()  # Re-display the original frame to select a new file2

money_type = "DOLLAR"

def toggle_currency():
    global money_type
    if money_type == "DOLLAR": money_type = "DIRHAM"; toggle_currency_button.configure(image=off_icon)
    else: money_type = "DOLLAR"; toggle_currency_button.configure(image=on_icon)



root = tk.Tk()
root.title("Json to Excel and Docx Converter")

# Load images for on and off states
on_image = Image.open("images/usa_flag.png")
on_image = on_image.resize((45, 30))
on_icon = ImageTk.PhotoImage(on_image)

off_image = Image.open("images/Flag-United-Arab-Emirates.png")
off_image = off_image.resize((45, 30))
off_icon = ImageTk.PhotoImage(off_image)

# Get the screen width and height
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# Set the window size and position to center it on the screen
window_width = 600
window_height = 300
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# Create a frame with padding
frame = ttk.Frame(root, padding=20)
frame.grid(row=0, column=0)

# Create a style for buttons
style = ttk.Style()
style.configure("TButton", padding=(10, 5), font=("Arial", 12))

# Create a label for the title
title_label = ttk.Label(frame, text="Json to Excel and Docx Converter", font=("Arial", 16, "bold"))
title_label.grid(row=0, column=0, pady=(0, 20), columnspan=2)

# Create a button to select the Excel file
select_excel_button = ttk.Button(frame, text="Select Json File", command=convert_json_to_docx)
select_excel_button.grid(row=1, column=0, pady=(0, 10), padx=10, sticky='w')

excel_file_label = ttk.Label(frame, text="", font=("Arial", 10))
excel_file_label.grid(row=2, column=0, pady=(10, 10), padx=10, sticky='w')

currency_var = tk.BooleanVar(value=True)  # True for DOLLAR, False for DIRHAM
toggle_currency_button = ttk.Checkbutton(frame, image=on_icon, variable=currency_var, command=toggle_currency,style="TButton", onvalue=True, offvalue=False)
toggle_currency_button.grid(row=3, column=0, pady=(0, 10), padx=10, sticky='w')

# Center all widgets both vertically and horizontally
frame.grid_rowconfigure(1, weight=1)
frame.grid_columnconfigure(0, weight=1)

selected_json_file = None

root.mainloop()
